"""
Create Module 4: Model with Results Word documents
- One for Karan Parekh (Hierarchical OLS Regression)
- One for Chau Lee (Machine Learning Model Comparison)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

FIGURES = r"C:\Users\karan\Downloads\ADTA5940-Capstone\figures"
OUT_DIR = r"C:\Users\karan\Downloads\ADTA5940-Capstone"


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading.append(shd)


def style_doc(doc):
    """Apply consistent styling to document."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    para = style.paragraph_format
    para.space_after = Pt(6)
    para.line_spacing = 1.15
    return doc


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_figure(doc, filename, caption, width=5.5):
    """Insert a figure with caption, centered."""
    path = os.path.join(FIGURES, filename)
    if not os.path.exists(path):
        p = doc.add_paragraph(f"[Figure not found: {filename}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = doc.styles['Normal']
    for run in cap.runs:
        run.font.size = Pt(10)
        run.font.italic = True


def make_table(doc, headers, rows, col_widths=None):
    """Create a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)

    return table


# =============================================================
# DOCUMENT 1: KARAN PAREKH — Hierarchical OLS Regression
# =============================================================
def create_karan_doc():
    doc = Document()
    style_doc(doc)

    # ----- Title Page -----
    for _ in range(4):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("Module 4: Model with Results")
    run.font.size = Pt(22)
    run.bold = True

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t2.add_run("Remote Work, Work-Life Balance, and Federal Employee Job Satisfaction:\nA Hierarchical Regression Analysis")
    run.font.size = Pt(14)

    for _ in range(2):
        doc.add_paragraph()

    info_lines = [
        "Karan Parekh",
        "ADTA 5940 — Analytics Capstone Experience",
        "Section 501 | Spring 2026",
        "University of North Texas",
        "Dr. Denise Philpot",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12)

    doc.add_page_break()

    # ----- 1. Research Question -----
    add_heading_styled(doc, "1. Research Question", level=1)

    doc.add_paragraph(
        "For this part of our project, I wanted to look at one main question: does telework "
        "actually make federal employees more satisfied with their jobs, or is there something "
        "else going on? During our earlier EDA work, Chau Lee and I noticed that employees "
        "who telework regularly tend to report higher satisfaction scores than people required "
        "to come into the office. But that does not necessarily mean telework itself is the "
        "reason — it could be that teleworkers also happen to have better supervisors, stronger "
        "engagement, or more work-life balance support."
    )

    doc.add_paragraph(
        "So the formal research question is:"
    )

    rq = doc.add_paragraph()
    rq.paragraph_format.left_indent = Inches(0.5)
    rq.paragraph_format.right_indent = Inches(0.5)
    run = rq.add_run(
        "To what extent does telework frequency predict overall job satisfaction among federal "
        "employees, after accounting for work-life balance perceptions, employee engagement, "
        "and demographic characteristics?"
    )
    run.italic = True

    doc.add_paragraph(
        "I set up three hypotheses to test:"
    )

    hyp_items = [
        "H1: Telework frequency is a significant predictor of job satisfaction. Employees who "
        "telework routinely will report higher satisfaction than those required to work on-site.",
        "H2: Work-life balance and employee engagement each explain meaningful additional "
        "variance in satisfaction beyond what telework alone can explain.",
        "H3: Once we control for work-life balance and engagement, the telework effect "
        "will shrink significantly — suggesting those attitudes partially explain the raw telework gap."
    ]
    for h in hyp_items:
        p = doc.add_paragraph(h, style='List Bullet')

    # ----- 2. Data Preparation -----
    add_heading_styled(doc, "2. Data Preparation", level=1)

    doc.add_paragraph(
        "The data comes from OPM's 2024 Federal Employee Viewpoint Survey (FEVS) Public Release "
        "Data File. The raw file has 674,207 rows and 96 columns. Since the PRDF stores everything "
        "as text (including the survey responses), the first step was converting Likert-scale items "
        'to numeric values. Responses coded as "X" (Not Applicable) or blanks were set to missing.'
    )

    doc.add_paragraph(
        "I built two composite indices following OPM's published methodology:"
    )

    comp_items = [
        "Employee Engagement Index (EEI): The mean of three OPM sub-indices — Intrinsic Work "
        "Experience (5 items), Supervisor (5 items), and Leaders Lead (5 items). This gives one "
        "overall engagement score per person on a 1-5 scale.",
        "Work-Life Balance Composite (WLB): The mean of three items — Q34 (peer support for WLB), "
        "Q49 (supervisor supports WLB), and Q63 (senior leaders support WLB programs)."
    ]
    for item in comp_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        "For the telework variable, I used Q91 which asks about telework participation level. "
        "OPM codes it as: 1 = Routine/Remote, 2 = Situational/Infrequent, 3 = Required in Office, "
        "4 = Chooses Not to Telework. I created dummy variables with Routine/Remote as the "
        "reference group since that is the baseline we want to compare against."
    )

    doc.add_paragraph(
        "Demographics were recoded from OPM's letter codes: DAGEGRP (A = Under 40, B = 40+), "
        "DSEX (A = Male, B = Female), DSUPER (A = Non-Supervisor, B = Supervisor), and DFEDTEN "
        "(A = 10 yrs or less, B = 11-20, C = 20+). Federal tenure was split into two dummies "
        "with 10-years-or-less as the reference."
    )

    doc.add_paragraph(
        "After listwise deletion of any rows with missing values on model variables, the final "
        "analytic sample was n = 517,697. This is a large sample, which means even tiny effects "
        "will be statistically significant — so throughout the results I focus on effect sizes "
        "and practical significance, not just p-values."
    )

    # ----- 3. Formal Model -----
    add_heading_styled(doc, "3. Formal Model", level=1)

    doc.add_paragraph(
        "I used a hierarchical (blockwise) OLS multiple regression. The idea is to enter "
        "predictors in stages so we can see how much each group of variables adds to the "
        "model. This lets us answer the question: what does telework explain on its own, "
        "and how much does that change when we add other factors?"
    )

    doc.add_paragraph("The four models are nested as follows:")

    model_specs = [
        ("Model 1 — Telework Only:", "Job Satisfaction = intercept + telework dummies"),
        ("Model 2 — Add Work-Life Balance:", "Job Satisfaction = intercept + telework + WLB composite"),
        ("Model 3 — Add Employee Engagement:", "Job Satisfaction = intercept + telework + WLB + EEI"),
        ("Model 4 — Full Model with Demographics:", "Job Satisfaction = intercept + telework + WLB + EEI + supervisor + gender + age + tenure"),
    ]
    for label, spec in model_specs:
        p = doc.add_paragraph()
        run = p.add_run(label + " ")
        run.bold = True
        p.add_run(spec)

    doc.add_paragraph(
        "The dependent variable is Q70: \"Considering everything, how satisfied are you with "
        "your job?\" on a 1-5 Likert scale. I am treating this as continuous, which is standard "
        "practice for Likert scales in social science research, though I acknowledge the ordinal "
        "nature is a limitation."
    )

    # Variable table
    add_heading_styled(doc, "3.1 Variable Definitions", level=2)

    var_headers = ["Variable", "Measurement", "Source"]
    var_rows = [
        ["Job Satisfaction (DV)", "Q70: 1-5 Likert (Very Dissatisfied to Very Satisfied)", "FEVS 2024"],
        ["TW: Routine/Remote", "Reference group (Q91 = 1)", "Q91"],
        ["TW: Infrequent", "= 1 if situational telework; 0 otherwise", "Q91 = 2"],
        ["TW: Required in Office", "= 1 if required on-site; 0 otherwise", "Q91 = 3"],
        ["TW: Chooses Not To", "= 1 if chooses not to telework; 0 otherwise", "Q91 = 4"],
        ["WLB Composite", "Mean of Q34, Q49, Q63 (1-5 scale)", "Computed"],
        ["EEI", "Mean of 3 OPM sub-indices (15 items total, 1-5)", "OPM formula"],
        ["Supervisor", "= 1 if supervisor/manager; 0 otherwise", "DSUPER"],
        ["Male", "= 1 if male; 0 if female", "DSEX"],
        ["Age 40+", "= 1 if 40 or older; 0 if under 40", "DAGEGRP"],
        ["Tenure 11-20 yrs", "= 1 if 11-20 years; 0 otherwise", "DFEDTEN"],
        ["Tenure 20+ yrs", "= 1 if 20+ years; 0 otherwise", "DFEDTEN"],
    ]
    make_table(doc, var_headers, var_rows)

    # ----- 4. Results -----
    add_heading_styled(doc, "4. Results", level=1)

    # 4.1 Descriptive
    add_heading_styled(doc, "4.1 Descriptive Overview", level=2)

    doc.add_paragraph(
        "Table 1 shows descriptive stats for the main continuous variables. The average "
        "job satisfaction is 3.82 out of 5 (SD = 1.09), which means most federal employees "
        "fall between \"Neither\" and \"Satisfied.\" Both WLB (mean = 4.02) and EEI (mean = 3.95) "
        "are above the midpoint, reflecting generally positive perceptions."
    )

    desc_headers = ["Variable", "Mean", "SD", "Min", "Max"]
    desc_rows = [
        ["Job Satisfaction (Q70)", "3.815", "1.088", "1.0", "5.0"],
        ["Work-Life Balance Composite", "4.018", "0.898", "1.0", "5.0"],
        ["Employee Engagement Index", "3.951", "0.817", "1.0", "5.0"],
        ["EEI — Intrinsic Work Experience", "3.941", "0.861", "1.0", "5.0"],
        ["EEI — Supervisor", "4.244", "0.931", "1.0", "5.0"],
        ["EEI — Leaders Lead", "3.669", "1.041", "1.0", "5.0"],
    ]
    p = doc.add_paragraph()
    run = p.add_run("Table 1. Descriptive Statistics (N = 646,545)")
    run.bold = True
    run.font.size = Pt(10)
    make_table(doc, desc_headers, desc_rows)

    doc.add_paragraph(
        "Looking at the telework distribution: 42.8% telework routinely or are fully remote, "
        "33.6% telework situationally, 19.7% are required in the office, and 3.8% choose not "
        "to telework. The raw satisfaction gap is clear — routine teleworkers average 3.94 "
        "while required-in-office employees average 3.54, a difference of 0.40 points. A "
        "one-way ANOVA confirmed this is statistically significant (F(3, 639,584) = 4,046.84, "
        "p < .001), though the effect size is small (eta-squared = 0.019)."
    )

    add_figure(doc, "fig1_satisfaction_by_telework.png",
               "Figure 1. Mean Job Satisfaction by Telework Category (error bars = 95% CI)", 5.0)

    # 4.2 Hierarchical regression
    add_heading_styled(doc, "4.2 Hierarchical Regression Results", level=2)

    doc.add_paragraph(
        "Table 2 shows the model fit at each step. This is really the core of the analysis "
        "and I think the pattern tells a clear story."
    )

    hier_headers = ["Metric", "Model 1", "Model 2", "Model 3", "Model 4"]
    hier_rows = [
        ["Block Added", "Telework", "+ WLB", "+ EEI", "+ Demographics"],
        ["R-squared", ".0200", ".4091", ".5945", ".5966"],
        ["Adj. R-squared", ".0200", ".4091", ".5945", ".5966"],
        ["Delta R-squared", "—", ".3891", ".1854", ".0021"],
        ["F-statistic", "3,516.10", "89,600.50", "151,801.68", "76,564.48"],
        ["Cohen's f-squared", "0.020 (small)", "0.692 (large)", "1.466 (large)", "1.479 (large)"],
    ]
    p = doc.add_paragraph()
    run = p.add_run("Table 2. Model Comparison — Hierarchical Regression (n = 517,697)")
    run.bold = True
    run.font.size = Pt(10)
    make_table(doc, hier_headers, hier_rows)

    doc.add_paragraph("")  # spacer

    doc.add_paragraph(
        "Here is what stands out to me from this progression:"
    )

    findings = [
        "Telework alone only explains 2% of the variance in satisfaction (Model 1). That is "
        "surprisingly low given all the attention telework gets in the news. The effect size "
        "is small (f-squared = 0.020).",
        "Adding Work-Life Balance causes R-squared to jump from .02 to .41, an increase of "
        "nearly 39 percentage points. This is by far the biggest single jump. Employees' "
        "perceptions of WLB support are way more predictive than where they physically sit.",
        "Adding Employee Engagement pushes R-squared to .595, another 18.5 percentage point "
        "increase. Together, WLB and EEI explain about 60% of the variation in satisfaction — "
        "that is a strong model for survey data.",
        "Demographics barely matter — they add only 0.2 percentage points. Age, gender, "
        "tenure, and supervisory status are statistically significant (because of the huge "
        "sample) but practically irrelevant."
    ]
    for f in findings:
        doc.add_paragraph(f, style='List Number')

    add_figure(doc, "fig2_r2_progression.png",
               "Figure 2. Cumulative R-squared by Model Block", 5.0)

    # 4.3 Full model coefficients
    add_heading_styled(doc, "4.3 Full Model Coefficients", level=2)

    doc.add_paragraph(
        "Table 3 shows the coefficients for the full model (Model 4). The standardized "
        "coefficients (Beta) let us compare the relative importance of each predictor."
    )

    coef_headers = ["Predictor", "B", "SE", "Beta", "p", "95% CI"]
    coef_rows = [
        ["Intercept", "-0.469", "0.006", "—", "< .001", "[-0.480, -0.458]"],
        ["TW: Infrequent", "-0.002", "0.002", "-0.001", ".326", "[-0.007, 0.002]"],
        ["TW: Required in Office", "0.114", "0.003", "0.043", "< .001", "[0.109, 0.119]"],
        ["TW: Chooses Not To", "0.090", "0.005", "0.016", "< .001", "[0.080, 0.099]"],
        ["WLB Composite", "0.111", "0.002", "0.091", "< .001", "[0.107, 0.114]"],
        ["Employee Engagement Index", "0.945", "0.002", "0.705", "< .001", "[0.941, 0.949]"],
        ["Supervisor", "-0.019", "0.003", "-0.007", "< .001", "[-0.024, -0.015]"],
        ["Male", "0.018", "0.002", "0.009", "< .001", "[0.015, 0.022]"],
        ["Age 40+", "0.088", "0.003", "0.036", "< .001", "[0.083, 0.093]"],
        ["Tenure 11-20 yrs", "-0.002", "0.002", "-0.001", ".536", "[-0.006, 0.003]"],
        ["Tenure 20+ yrs", "0.049", "0.003", "0.020", "< .001", "[0.044, 0.055]"],
    ]
    p = doc.add_paragraph()
    run = p.add_run("Table 3. OLS Coefficients — Full Model (Model 4)")
    run.bold = True
    run.font.size = Pt(10)
    make_table(doc, coef_headers, coef_rows)

    doc.add_paragraph(
        "R-squared = .5966, F(10, 517686) = 76,564.48, p < .001"
    ).runs[0].font.size = Pt(10)

    doc.add_paragraph(
        "The biggest takeaway from this table is that Employee Engagement (Beta = 0.705) "
        "absolutely dominates. A one-unit increase in EEI is associated with a 0.945-point "
        "increase in satisfaction — nearly a full point on the 5-point scale. Work-Life Balance "
        "is second most important (Beta = 0.091), and everything else is quite small in "
        "comparison."
    )

    add_figure(doc, "fig3_coefficients_forest.png",
               "Figure 3. Standardized Regression Coefficients — Full Model", 5.5)

    # 4.4 Telework coefficient progression
    add_heading_styled(doc, "4.4 The Telework Reversal (Simpson's Paradox)", level=2)

    doc.add_paragraph(
        "This was probably the most interesting finding in the analysis. Table 4 shows how "
        "the telework coefficients change as we add controls."
    )

    tw_headers = ["Predictor", "Model 1", "Model 2", "Model 3", "Model 4"]
    tw_rows = [
        ["TW: Infrequent", "-0.121***", "0.013***", "-0.002", "-0.002"],
        ["TW: Required in Office", "-0.407***", "+0.094***", "+0.109***", "+0.114***"],
        ["TW: Chooses Not To", "-0.029***", "+0.127***", "+0.096***", "+0.089***"],
    ]
    p = doc.add_paragraph()
    run = p.add_run("Table 4. Telework Coefficients Across Models")
    run.bold = True
    run.font.size = Pt(10)
    make_table(doc, tw_headers, tw_rows)

    doc.add_paragraph(
        "Look at the \"Required in Office\" row. In Model 1, the coefficient is -0.407, "
        "meaning office-required employees score nearly half a point lower than teleworkers. "
        "But once we add Work-Life Balance in Model 2, the sign flips to +0.094. By the full "
        "model, it is +0.114."
    )

    doc.add_paragraph(
        "This is a textbook case of Simpson's Paradox — the relationship between two variables "
        "(telework and satisfaction) reverses when you account for a confounding variable "
        "(WLB and engagement). What is really happening is that in-office employees tend to "
        "have lower WLB support and lower engagement scores, and those attitudinal factors "
        "are what actually drive their lower satisfaction. Once we equalize WLB and engagement "
        "across groups, office-required employees are actually slightly more satisfied. This "
        "could reflect things like self-selection — people in roles they chose to do in person, "
        "or jobs with inherent in-person rewards."
    )

    # 4.5 WLB scatter
    add_figure(doc, "fig4_wlb_vs_satisfaction.png",
               "Figure 4. Work-Life Balance vs. Job Satisfaction (r = .64)", 4.5)

    # 4.6 Diagnostics
    add_heading_styled(doc, "4.5 Model Diagnostics", level=2)

    diag_items = [
        "Multicollinearity: VIF was below 5 for all telework and demographic variables. "
        "However, WLB (VIF = 56.7) and EEI (VIF = 58.7) are highly collinear with each other, "
        "which makes sense since they measure related concepts. This does not affect the "
        "overall R-squared but does mean we should be cautious interpreting their individual "
        "coefficients in isolation.",
        "Autocorrelation: Durbin-Watson = 1.989, which is very close to 2 (no autocorrelation). "
        "Expected for cross-sectional survey data.",
        "Heteroscedasticity: The Breusch-Pagan test was significant (p < .001), so I also "
        "computed HC3 robust standard errors. All significance conclusions stayed the same.",
        "Residual normality: Slight negative skew (-0.531) and mild leptokurtosis (1.757). "
        "With over 500,000 observations, the Central Limit Theorem makes this a non-issue "
        "for inference.",
    ]
    for item in diag_items:
        doc.add_paragraph(item, style='List Bullet')

    add_figure(doc, "fig6_diagnostics.png",
               "Figure 6. Residual Diagnostics — Residuals vs. Fitted (left) and Distribution (right)", 5.5)

    # ----- 5. Discussion -----
    add_heading_styled(doc, "5. Discussion and Conclusions", level=1)

    doc.add_paragraph(
        "Going back to my three hypotheses:"
    )

    h_results = [
        "H1 — Partially supported. Telework is statistically significant but practically weak. "
        "By itself it only explains 2% of variance, and the coefficients reverse once we add "
        "controls. Telework frequency alone is not a reliable predictor of satisfaction.",
        "H2 — Strongly supported. WLB added 38.9 percentage points to R-squared and EEI "
        "added another 18.5 points. These are large, meaningful effects. Together they get "
        "the model to nearly 60% variance explained.",
        "H3 — Supported. The telework effect shrank dramatically and even reversed direction "
        "after controlling for WLB and engagement. This confirms that the raw telework-satisfaction "
        "gap is driven by confounding — teleworkers happen to have better engagement and "
        "WLB support, not because telework itself causes satisfaction.",
    ]
    for h in h_results:
        doc.add_paragraph(h, style='List Bullet')

    add_heading_styled(doc, "5.1 What This Means for Federal Agencies", level=2)

    doc.add_paragraph(
        "The practical implication is that agencies should not treat expanding telework as "
        "a silver bullet for improving satisfaction. The data shows that what actually moves "
        "the needle is employee engagement and work-life balance support. Agencies that invest "
        "in strong supervision, meaningful work experiences, and genuine support for work-life "
        "balance will see satisfied employees regardless of where those employees physically sit."
    )

    doc.add_paragraph(
        "This does not mean telework is unimportant — there are many good reasons to offer it "
        "(recruiting, retention, cost savings). But the expectation that moving people remote "
        "will automatically increase satisfaction is not supported by this data."
    )

    add_heading_styled(doc, "5.2 Limitations", level=2)

    limits = [
        "Cross-sectional design: This is one snapshot in time, so I cannot make causal claims. "
        "The mediation pattern is consistent with telework affecting engagement which affects "
        "satisfaction, but it could go the other way.",
        "High VIF for WLB and EEI: These two composites are strongly correlated. In a future "
        "analysis, structural equation modeling could better separate their individual contributions.",
        "Self-selection: People who telework may be different in unobserved ways (higher grade, "
        "more autonomy in their role) that independently predict satisfaction.",
        "Likert as continuous: Treating 1-5 ordinal data as continuous in OLS is common but "
        "imperfect. An ordinal logistic regression could serve as a robustness check.",
    ]
    for lim in limits:
        doc.add_paragraph(lim, style='List Bullet')

    # References
    add_heading_styled(doc, "References", level=1)

    refs = [
        "Bakker, A. B., & Demerouti, E. (2017). Job Demands-Resources theory: Taking stock "
        "and looking forward. Journal of Occupational Health Psychology, 22(3), 273-285.",
        "Bloom, N., Liang, J., Roberts, J., & Ying, Z. J. (2015). Does working from home "
        "work? Evidence from a Chinese experiment. Quarterly Journal of Economics, 130(1), 165-218.",
        "Caillier, J. G. (2012). The impact of teleworking on work motivation in a U.S. "
        "federal government agency. American Review of Public Administration, 42(4), 461-480.",
        "Cohen, J. (1988). Statistical power analysis for the behavioral sciences (2nd ed.). Erlbaum.",
        "De Vries, H., Tummers, L., & Bekkers, V. (2019). The benefits of teleworking in the "
        "public sector: Reality or rhetoric? Review of Public Personnel Administration, 39(4), 570-593.",
        "Office of Personnel Management. (2024). Federal Employee Viewpoint Survey: Public "
        "Release Data File and Codebook. Washington, DC: OPM.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        for run in p.runs:
            run.font.size = Pt(11)

    # Save
    path = os.path.join(OUT_DIR, "Module4_Model_Results_Karan_Parekh.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


# =============================================================
# DOCUMENT 2: CHAU LEE — Machine Learning Model Comparison
# =============================================================
def create_chau_doc():
    doc = Document()
    style_doc(doc)

    # ----- Title Page -----
    for _ in range(4):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("Module 4: Model with Results")
    run.font.size = Pt(22)
    run.bold = True

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t2.add_run("Predicting Federal Employee Job Satisfaction:\nA Machine Learning Model Comparison")
    run.font.size = Pt(14)

    for _ in range(2):
        doc.add_paragraph()

    info_lines = [
        "Chau Lee",
        "ADTA 5940 — Analytics Capstone Experience",
        "Section 501 | Spring 2026",
        "University of North Texas",
        "Dr. Denise Philpot",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12)

    doc.add_page_break()

    # ----- 1. Research Question -----
    add_heading_styled(doc, "1. Research Question", level=1)

    doc.add_paragraph(
        "While Karan's analysis focused on explaining which factors drive satisfaction "
        "(an inferential question), my part of the project asks a different but complementary "
        "question: can we actually predict how satisfied a federal employee will be? And if so, "
        "which algorithm does it best?"
    )

    doc.add_paragraph("The formal research question is:")

    rq = doc.add_paragraph()
    rq.paragraph_format.left_indent = Inches(0.5)
    rq.paragraph_format.right_indent = Inches(0.5)
    run = rq.add_run(
        "Which machine learning model most accurately predicts overall job satisfaction "
        "among federal employees, given their telework frequency, work-life balance perception, "
        "employee engagement, and demographic characteristics?"
    )
    run.italic = True

    doc.add_paragraph(
        "This is a pure prediction problem — I am not trying to explain why someone is "
        "satisfied, but rather how well we can predict their satisfaction score from features "
        "that are available in the FEVS data. This has practical value because if an agency "
        "can predict which employees are at risk for low satisfaction, they can intervene early."
    )

    doc.add_paragraph("The hypotheses guiding this work:")

    hyp_items = [
        "H1: Ensemble tree-based models (Random Forest, Gradient Boosting, XGBoost, LightGBM) "
        "will outperform simpler models like Linear Regression because the relationship between "
        "features and satisfaction is likely nonlinear.",
        "H2: The Employee Engagement Index will be the most important feature for predicting "
        "satisfaction across model types.",
        "H3: The best model will achieve meaningfully higher accuracy than a basic linear "
        "regression baseline, justifying the added complexity.",
    ]
    for h in hyp_items:
        doc.add_paragraph(h, style='List Bullet')

    # ----- 2. Data Preparation -----
    add_heading_styled(doc, "2. Data Preparation", level=1)

    doc.add_paragraph(
        "I used the same 2024 FEVS Public Release Data File (674,207 rows, 96 columns) as "
        "Karan's analysis, but my feature engineering was a bit different since the goal here "
        "is prediction rather than inference."
    )

    doc.add_paragraph(
        "The target variable is Q69 (\"Considering everything, how satisfied are you with your "
        "job?\") on a 1-5 scale. The features I selected are:"
    )

    feat_items = [
        "Q61 — Telework frequency (1-5 scale: never to 3+ days/week)",
        "Q64 — Work-life balance satisfaction (1-5 Likert)",
        "EEI — Employee Engagement Index, calculated as the mean of 7 OPM items "
        "(Q3, Q4, Q6, Q11, Q12, Q13, Q14). I collapsed the 7 items into one composite to "
        "reduce multicollinearity and follow OPM's established methodology.",
        "DAGEGRP — Age group (0 = Under 40, 1 = 40+)",
        "DSEX — Gender (0/1 binary)",
        "DSUPER — Supervisory status (0 = Non-supervisor, 1 = Supervisor)",
        "DFEDTEN — Federal tenure (0 = 10 yrs or less, 1 = 11-20 yrs, 2 = 20+ yrs)",
    ]
    for f in feat_items:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_paragraph(
        "Survey responses coded as \"X\" (Not Applicable) were treated as missing and converted "
        "to NaN. Demographic letter codes (A, B, C) were mapped to numeric values. After "
        "dropping all rows with any missing data (listwise deletion), the cleaned dataset has "
        "453,082 rows — about 67% of the original data. The 33% drop is mainly due to respondents "
        "who skipped one or more survey items."
    )

    doc.add_paragraph(
        "The data was split 80/20 into training (362,465 samples) and test (90,617 samples) sets. "
        "All features were standardized using StandardScaler (fit on training, applied to test) "
        "to ensure fair comparison across algorithms with different scale sensitivities. For the "
        "two most computationally expensive models — KNN and MLP Neural Net — I used a random "
        "subsample of 30,000 training rows to keep runtime manageable."
    )

    # ----- 3. Model Selection Process -----
    add_heading_styled(doc, "3. Model Selection Process", level=1)

    doc.add_paragraph(
        "I tested 14 different regression algorithms spanning four categories. The idea was "
        "to cast a wide net and let the data tell us which approach works best, rather than "
        "picking one model upfront."
    )

    cat_items = [
        "Linear models (4): Linear Regression, Ridge, Lasso, Elastic Net — these are our "
        "baselines. If a complex model cannot beat basic linear regression, there is no "
        "justification for the added complexity.",
        "Tree-based models (3): Decision Tree, Random Forest, Extra Trees — these can capture "
        "nonlinear patterns and feature interactions without manual specification.",
        "Boosting models (4): Gradient Boosting, AdaBoost, XGBoost, LightGBM — these build "
        "trees sequentially, each one learning from the previous tree's mistakes. They "
        "typically achieve the best performance on tabular data.",
        "Other models (3): Bagging, KNN (K-Nearest Neighbors), MLP Neural Net — included "
        "for completeness and comparison.",
    ]
    for item in cat_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        "Each model was evaluated using the same protocol: train on the training set, "
        "predict on the held-out test set, and compute R-squared, MAE (Mean Absolute Error), "
        "and RMSE (Root Mean Squared Error). Additionally, I ran 5-fold cross-validation "
        "on a 50,000-row subsample to check for stability and overfitting."
    )

    add_heading_styled(doc, "3.1 Evaluation Metrics", level=2)

    metric_items = [
        "R-squared: The proportion of variance in satisfaction explained by the model. "
        "Higher is better. A value of 0.53 means the model explains 53% of the variation.",
        "MAE: The average absolute error in predictions, measured in the same units as the "
        "target variable. An MAE of 0.61 means predictions are off by about 0.61 points on "
        "the 1-5 scale, on average.",
        "RMSE: Similar to MAE but penalizes larger errors more heavily. Useful for identifying "
        "models that make occasional big mistakes.",
        "Cross-validation R-squared: The average R-squared across 5 different train/test splits. "
        "A model with high test R-squared but low CV R-squared might be overfitting.",
    ]
    for m in metric_items:
        doc.add_paragraph(m, style='List Bullet')

    # ----- 4. Results -----
    add_heading_styled(doc, "4. Results", level=1)

    # 4.1 Correlation analysis
    add_heading_styled(doc, "4.1 Feature Correlations", level=2)

    doc.add_paragraph(
        "Before running the models, I looked at how the features correlate with each other "
        "and with the target variable. Figure 1 shows the correlation matrix."
    )

    add_figure(doc, "correlation_matrix.png",
               "Figure 1. Correlation Matrix — Features vs. Job Satisfaction (Q69)", 4.5)

    doc.add_paragraph(
        "The three strongest predictors are EEI (r = 0.65), Q64 work-life balance (r = 0.63), "
        "and Q61 telework (r = 0.60). The four demographic features all have very weak "
        "correlations with satisfaction (r < 0.06). This gives us a preview of what the "
        "models will find — the attitudinal measures carry most of the predictive signal."
    )

    # 4.2 Model comparison
    add_heading_styled(doc, "4.2 Model Comparison Results", level=2)

    doc.add_paragraph(
        "Table 1 presents the results for all 14 models, ranked by test R-squared."
    )

    model_headers = ["Rank", "Model", "R² (Test)", "MAE", "RMSE", "CV R² Mean", "Time (s)"]
    model_rows = [
        ["1", "XGBoost", "0.5322", "0.6056", "0.8016", "0.5241", "2.69"],
        ["2", "LightGBM", "0.5322", "0.6057", "0.8015", "0.5243", "0.38"],
        ["3", "Gradient Boosting", "0.5321", "0.6057", "0.8016", "0.5239", "18.05"],
        ["4", "Extra Trees", "0.5308", "0.6074", "0.8027", "0.5230", "2.27"],
        ["5", "Random Forest", "0.5307", "0.6063", "0.8028", "0.5142", "3.13"],
        ["6", "Decision Tree", "0.5287", "0.6072", "0.8045", "0.5079", "0.21"],
        ["7", "MLP Neural Net", "0.5259", "0.6202", "0.8070", "0.5268", "2.42"],
        ["8", "Linear Regression", "0.5254", "0.6172", "0.8073", "0.5242", "0.06"],
        ["9", "Ridge Regression", "0.5254", "0.6172", "0.8073", "0.5230", "0.05"],
        ["10", "Elastic Net", "0.5253", "0.6184", "0.8075", "0.5280", "0.04"],
        ["11", "Lasso Regression", "0.5250", "0.6192", "0.8077", "0.5194", "0.06"],
        ["12", "Bagging", "0.5138", "0.6136", "0.8172", "0.4706", "2.44"],
        ["13", "AdaBoost", "0.4703", "0.6945", "0.8530", "0.4727", "2.45"],
        ["14", "KNN", "0.4689", "0.6425", "0.8540", "0.4657", "0.03"],
    ]
    p = doc.add_paragraph()
    run = p.add_run("Table 1. Model Comparison Results — All 14 Models")
    run.bold = True
    run.font.size = Pt(10)
    make_table(doc, model_headers, model_rows)

    doc.add_paragraph("")

    doc.add_paragraph(
        "A few observations from the results:"
    )

    obs_items = [
        "The top 3 models (XGBoost, LightGBM, Gradient Boosting) are essentially tied at "
        "R-squared = 0.532. The differences between them are tiny — less than 0.0002.",
        "The boosting models do outperform linear regression, but only by about 0.7 percentage "
        "points (0.5322 vs. 0.5254). That is a real improvement but smaller than I expected. "
        "This suggests the relationship between features and satisfaction is mostly linear, "
        "with some nonlinear patterns that boosting can pick up.",
        "KNN and AdaBoost performed noticeably worse than everything else. KNN likely suffers "
        "because it was trained on a smaller subsample, and the curse of dimensionality "
        "makes distance-based methods less effective even with only 7 features.",
        "LightGBM trained in just 0.38 seconds compared to 18 seconds for Gradient Boosting "
        "and 2.7 seconds for XGBoost — while achieving the same accuracy. This speed advantage "
        "matters for deployment and iterative analysis.",
    ]
    for obs in obs_items:
        doc.add_paragraph(obs, style='List Number')

    add_figure(doc, "model_r2_comparison.png",
               "Figure 2. R-squared Comparison Across All 14 Models", 5.5)

    # 4.3 Statistical significance tests
    add_heading_styled(doc, "4.3 Are the Top Models Really Different?", level=2)

    doc.add_paragraph(
        "Since the top models are so close in performance, I ran paired t-tests on their "
        "5-fold cross-validation scores to check if the differences are statistically significant."
    )

    sig_headers = ["Comparison", "Mean R² A", "Mean R² B", "Diff", "p-value", "Significant?"]
    sig_rows = [
        ["XGBoost vs. LightGBM", "0.5241", "0.5243", "-0.0002", "0.944", "No"],
        ["XGBoost vs. Gradient Boosting", "0.5241", "0.5239", "0.0002", "0.968", "No"],
        ["XGBoost vs. Extra Trees", "0.5241", "0.5230", "0.0011", "0.874", "No"],
        ["XGBoost vs. Random Forest", "0.5241", "0.5142", "0.0099", "0.027", "Yes"],
        ["LightGBM vs. Gradient Boosting", "0.5243", "0.5239", "0.0004", "0.953", "No"],
    ]
    p = doc.add_paragraph()
    run = p.add_run("Table 2. Paired t-tests on Cross-Validation Scores (Top 5 Models)")
    run.bold = True
    run.font.size = Pt(10)
    make_table(doc, sig_headers, sig_rows)

    doc.add_paragraph(
        "The top 4 models (XGBoost, LightGBM, Gradient Boosting, Extra Trees) show no "
        "statistically significant differences from each other. Only Random Forest is "
        "significantly worse than XGBoost (p = 0.027). This means we have some freedom "
        "in choosing our final model — accuracy alone does not distinguish them."
    )

    add_figure(doc, "cv_boxplots.png",
               "Figure 3. Cross-Validation R-squared Distribution by Model", 5.5)

    # 4.4 Model selection
    add_heading_styled(doc, "4.4 Why LightGBM?", level=2)

    doc.add_paragraph(
        "Since the top models perform the same statistically, I selected LightGBM as the "
        "primary model based on secondary criteria:"
    )

    lgb_reasons = [
        "Speed: LightGBM trained in 0.38 seconds — 7 times faster than XGBoost (2.69s) and "
        "47 times faster than Gradient Boosting (18.05s). For an agency that might retrain "
        "this model annually on new FEVS data, speed matters.",
        "Highest CV R-squared: LightGBM had the highest cross-validation mean (0.5243), "
        "suggesting it generalizes slightly better to unseen data.",
        "Best RMSE: 0.8015, meaning it makes fewer large prediction errors compared to "
        "the other top models.",
        "Differentiation: Karan used XGBoost in a previous course project, and since LightGBM "
        "performs identically here, it demonstrates familiarity with multiple tools.",
    ]
    for r in lgb_reasons:
        doc.add_paragraph(r, style='List Bullet')

    # 4.5 Feature importance
    add_heading_styled(doc, "4.5 Feature Importance", level=2)

    doc.add_paragraph(
        "One advantage of tree-based models is that they tell us which features they relied "
        "on most for predictions. Table 3 shows the LightGBM feature importances."
    )

    fi_headers = ["Feature", "Importance (%)", "Description"]
    fi_rows = [
        ["EEI", "34.4%", "Employee Engagement Index (7-item composite)"],
        ["Q61", "21.2%", "Telework frequency"],
        ["Q64", "18.9%", "Work-life balance satisfaction"],
        ["DFEDTEN", "10.4%", "Federal tenure"],
        ["DSUPER", "5.8%", "Supervisory status"],
        ["DSEX", "4.9%", "Gender"],
        ["DAGEGRP", "4.5%", "Age group"],
    ]
    p = doc.add_paragraph()
    run = p.add_run("Table 3. LightGBM Feature Importances")
    run.bold = True
    run.font.size = Pt(10)
    make_table(doc, fi_headers, fi_rows)

    doc.add_paragraph(
        "The top 3 features — EEI, Q61, and Q64 — account for 74.5% of the model's total "
        "importance. This aligns with the correlation analysis and with Karan's regression "
        "finding that engagement and WLB perceptions are the dominant predictors. Demographics "
        "contribute about 25.6% combined, with federal tenure being the most noteworthy "
        "demographic feature at 10.4%."
    )

    add_figure(doc, "feature_importances.png",
               "Figure 4. Feature Importances — LightGBM", 5.0)

    # 4.6 Prediction quality
    add_heading_styled(doc, "4.6 Prediction Quality", level=2)

    doc.add_paragraph(
        "Figure 5 shows actual versus predicted satisfaction scores. If the model were "
        "perfect, all points would fall on the diagonal line. The predictions cluster "
        "reasonably well around the diagonal, especially in the middle range (scores 3-4 "
        "where most employees fall). The model struggles more at the extremes — it tends "
        "to underpredict for very satisfied employees (5s) and overpredict for very "
        "dissatisfied employees (1s). This is common in regression on bounded scales."
    )

    add_figure(doc, "actual_vs_predicted.png",
               "Figure 5. Actual vs. Predicted Job Satisfaction — LightGBM", 4.5)

    doc.add_paragraph(
        "In practical terms, with an MAE of 0.61, the model's predictions are off by about "
        "half a point on average. On a 1-5 scale, that is decent — it means if someone's "
        "true satisfaction is 4.0, the model would typically predict somewhere between 3.4 "
        "and 4.6. Not precise enough for individual-level decisions, but useful for "
        "identifying groups or units that are at risk."
    )

    add_figure(doc, "final_ranking.png",
               "Figure 6. Final Model Ranking — All 14 Models", 5.5)

    # ----- 5. Discussion -----
    add_heading_styled(doc, "5. Discussion and Conclusions", level=1)

    doc.add_paragraph("Revisiting my hypotheses:")

    h_results = [
        "H1 — Partially supported. The boosting models did outperform linear regression, "
        "but only by about 0.7 percentage points. The advantage is real but modest, suggesting "
        "the underlying relationship is mostly linear with minor nonlinear patterns.",
        "H2 — Supported. EEI is the single most important feature at 34.4% importance, "
        "followed by telework (21.2%) and work-life balance (18.9%). This matches the "
        "theoretical expectation and Karan's regression findings.",
        "H3 — Partially supported. The best model does beat Linear Regression, and the "
        "improvement is consistent across all metrics (R-squared, MAE, RMSE). But the gap "
        "is small enough that for many practical purposes, a simpler model might suffice.",
    ]
    for h in h_results:
        doc.add_paragraph(h, style='List Bullet')

    add_heading_styled(doc, "5.1 Practical Applications", level=2)

    doc.add_paragraph(
        "A model like this could be used by OPM or individual agencies to flag units or "
        "offices where predicted satisfaction is low, triggering early intervention. "
        "For example, if a regional office has features suggesting low engagement and limited "
        "telework, the model would predict lower satisfaction for those employees. HR could "
        "proactively address engagement before satisfaction drops further."
    )

    doc.add_paragraph(
        "The feature importance results also have a clear message for policy: if agencies "
        "want to improve the prediction (and by extension, the reality) of job satisfaction, "
        "the highest-return investments are in employee engagement programs and work-life "
        "balance support, not just telework expansion."
    )

    add_heading_styled(doc, "5.2 Limitations", level=2)

    limits = [
        "Moderate R-squared of 0.53: The model explains about half the variance. The other "
        "half likely comes from factors not in the FEVS PRDF — things like pay, specific "
        "agency culture, individual personality, or recent life events.",
        "Limited feature set: With only 7 features (plus Q69 as target), there is a ceiling "
        "on how accurate the model can get. The full FEVS has 96 columns; using more could "
        "improve predictions but would require more careful feature selection.",
        "Subsample for slow models: KNN and MLP were only trained on 30,000 rows instead "
        "of the full 362,000. Their rankings might improve with full training data, though "
        "likely not enough to beat the boosting models.",
        "No hyperparameter optimization for the final selection: I ran a brief tuning "
        "exercise and found that the default parameters already performed near-optimally. "
        "More extensive tuning (e.g., Bayesian optimization) might yield marginal gains.",
    ]
    for lim in limits:
        doc.add_paragraph(lim, style='List Bullet')

    # References
    add_heading_styled(doc, "References", level=1)

    refs = [
        "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5-32.",
        "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. "
        "Proceedings of the 22nd ACM SIGKDD International Conference, 785-794.",
        "Friedman, J. H. (2001). Greedy function approximation: A gradient boosting "
        "machine. The Annals of Statistics, 29(5), 1189-1232.",
        "Ke, G., Meng, Q., Finley, T., et al. (2017). LightGBM: A highly efficient "
        "gradient boosting decision tree. Advances in Neural Information Processing "
        "Systems, 30, 3146-3154.",
        "Office of Personnel Management. (2024). Federal Employee Viewpoint Survey: "
        "Public Release Data File and Codebook. Washington, DC: OPM.",
        "Office of Personnel Management. (2024). 2024 OPM FEVS Indices and Dimensions "
        "Guide. Washington, DC: OPM.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        for run in p.runs:
            run.font.size = Pt(11)

    # Save
    path = os.path.join(OUT_DIR, "Module4_Model_Results_Chau_Lee.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


# =============================================================
# MAIN
# =============================================================
if __name__ == "__main__":
    print("Creating Module 4 documents...\n")
    p1 = create_karan_doc()
    print()
    p2 = create_chau_doc()
    print(f"\nDone! Both documents saved to:\n  {p1}\n  {p2}")
