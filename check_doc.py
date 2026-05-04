from docx import Document

doc = Document('Module3_Scholarly_Review_Parekh_Lee_.docx')

# Check formatting
print("=== FONT CHECKS ===")
for i, para in enumerate(doc.paragraphs):
    if para.runs:
        r = para.runs[0]
        font = r.font
        if font.name and font.name != 'Times New Roman' and para.text.strip():
            print(f'P{i}: Font={font.name} (expected TNR)')
        if font.size and font.size != 152400 and para.text.strip():
            print(f'P{i}: Size={font.size} (expected 152400/12pt)')

# Check headers
print("\n=== HEADER/MARGIN CHECKS ===")
for section in doc.sections:
    header = section.header
    if header and header.paragraphs:
        for p in header.paragraphs:
            if p.text.strip():
                print(f'Header: "{p.text}" align={p.alignment}')
    print(f'Margins: top={section.top_margin}, bottom={section.bottom_margin}, left={section.left_margin}, right={section.right_margin}')

# Check line spacing on body paragraphs
print("\n=== LINE SPACING CHECK ===")
for i, para in enumerate(doc.paragraphs):
    pf = para.paragraph_format
    if pf.line_spacing and para.text.strip() and i > 14:
        print(f'P{i}: line_spacing={pf.line_spacing}')
        break

# Check for AI-telltale phrases
print("\n=== AI PHRASE SCAN ===")
WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
def get_full_text(para):
    parts = []
    tag_t = '{' + WNS + '}t'
    for child in para._element.iter():
        if child.tag == tag_t:
            parts.append(child.text or '')
    return ''.join(parts)

ai_phrases = [
    'delve', 'tapestry', 'landscape of', 'it is worth noting',
    'in conclusion', 'importantly,', 'it should be noted',
    'crucial to note', 'bears mentioning', 'multifaceted',
    'paradigm shift', 'holistic approach', 'synergy',
    'leveraging', 'in today\'s', 'ever-evolving',
    'game-changer', 'unlock', 'empower', 'robust',
    'comprehensive overview', 'serves as a testament',
    'stands as', 'plays a pivotal', 'groundbreaking'
]

full_text = ""
for para in doc.paragraphs:
    full_text += get_full_text(para) + " "

found_any = False
for phrase in ai_phrases:
    if phrase.lower() in full_text.lower():
        # find context
        idx = full_text.lower().find(phrase.lower())
        context = full_text[max(0,idx-30):idx+len(phrase)+30]
        print(f'  Found "{phrase}": ...{context}...')
        found_any = True

if not found_any:
    print("  No common AI phrases detected.")

# Check for missing author names (sentences starting with lowercase after period)
print("\n=== SENTENCE STRUCTURE CHECK ===")
import re
issues = []
for i, para in enumerate(doc.paragraphs):
    text = get_full_text(para)
    # Check for sentences that start with a verb (missing subject)
    patterns = [
        r'\.\s+(found|studied|conducted|established|made|argued|extended|showed|documented|defined)',
        r'^\s*(found|studied|conducted|established|made|argued|extended|showed|documented|defined)',
        r',\s+(found|studied|conducted|established|made|argued|extended|showed|documented|defined)',
    ]
    for pat in patterns:
        matches = re.findall(pat, text, re.IGNORECASE)
        if matches:
            for m in matches:
                idx = text.lower().find(m.lower())
                ctx = text[max(0,idx-40):idx+40]
                issues.append(f'P{i}: potential missing subject near "{m}": ...{ctx}...')

for issue in issues:
    print(f'  {issue}')

if not issues:
    print("  No missing subjects detected.")

# Check document metadata for AI traces
print("\n=== METADATA CHECK ===")
props = doc.core_properties
print(f'  Author: {props.author}')
print(f'  Last modified by: {props.last_modified_by}')
print(f'  Title: {props.title}')
print(f'  Comments: {props.comments}')
print(f'  Category: {props.category}')

print("\n=== DONE ===")
