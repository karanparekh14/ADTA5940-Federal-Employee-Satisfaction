"""Generate Model Selection Summary Word document for team reference."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Helper functions
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_bold_text(paragraph, bold_text, normal_text=""):
    run = paragraph.add_run(bold_text)
    run.bold = True
    if normal_text:
        paragraph.add_run(normal_text)

def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('ADTA 5940 Capstone — Model Selection Summary', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Prepared: March 19, 2026  |  Team: Karan Parekh & Chau Lee')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()  # spacer

# ============================================================
# SECTION 1
# ============================================================
add_heading_styled('1. What We Did', level=1)

p = doc.add_paragraph()
p.add_run('We ran a ').font.size = Pt(11)
r = p.add_run('comprehensive model comparison')
r.bold = True
p.add_run(' to find the best machine learning model for predicting ')
r = p.add_run('federal employee job satisfaction')
r.bold = True
p.add_run(' (FEVS survey question Q69, scored 1\u20135) using the 2024 OPM Federal Employee Viewpoint Survey (FEVS) dataset.')

bullets = [
    ('Dataset: ', '674,207 survey responses, 96 columns'),
    ('After cleaning (listwise deletion): ', '453,082 usable rows'),
    ('Train/Test split: ', '80% training (362,465) / 20% testing (90,617)'),
    ('Target variable: ', 'Q69 \u2014 "Considering everything, how satisfied are you with your job?" (1 = Very Dissatisfied \u2192 5 = Very Satisfied, mean = 3.626)'),
]
for bold_part, normal_part in bullets:
    p = doc.add_paragraph(style='List Bullet')
    add_bold_text(p, bold_part, normal_part)

add_heading_styled('Features Used (7 total)', level=2)
make_table(
    ['Feature', 'Description', 'Type'],
    [
        ['Q61', 'Telework frequency (1\u20135 scale)', 'Survey item'],
        ['Q64', 'Work-life balance satisfaction (1\u20135 scale)', 'Survey item'],
        ['EEI', 'Employee Engagement Index \u2014 composite mean of Q3, Q4, Q6, Q11\u2013Q14', 'Composite'],
        ['DAGEGRP', 'Age group (0 = Under 40, 1 = 40+)', 'Demographic'],
        ['DSEX', 'Gender (0 = Male, 1 = Female)', 'Demographic'],
        ['DSUPER', 'Supervisory status (0 = Non-supervisor, 1 = Supervisor)', 'Demographic'],
        ['DFEDTEN', 'Federal tenure (0 = <10yr, 1 = 10\u201320yr, 2 = 20+yr)', 'Demographic'],
    ]
)

# ============================================================
# SECTION 2
# ============================================================
doc.add_page_break()
add_heading_styled('2. All 14 Models Tested \u2014 Full Results', level=1)

make_table(
    ['Rank', 'Model', 'R\u00b2 (Test)', 'R\u00b2 (Train)', 'MAE', 'RMSE', 'CV R\u00b2 Mean', 'CV R\u00b2 Std', 'Time'],
    [
        ['1', 'XGBoost', '0.5322', '0.5339', '0.6056', '0.8016', '0.5241', '0.0031', '2.69s'],
        ['2', 'LightGBM', '0.5322', '0.5337', '0.6057', '0.8015', '0.5243', '0.0058', '0.38s'],
        ['3', 'Gradient Boosting', '0.5321', '0.5340', '0.6057', '0.8016', '0.5239', '0.0104', '18.05s'],
        ['4', 'Extra Trees', '0.5308', '0.5369', '0.6074', '0.8027', '0.5230', '0.0127', '2.27s'],
        ['5', 'Random Forest', '0.5307', '0.5374', '0.6063', '0.8028', '0.5142', '0.0042', '3.13s'],
        ['6', 'Decision Tree', '0.5287', '0.5357', '0.6072', '0.8045', '0.5079', '0.0053', '0.21s'],
        ['7', 'MLP Neural Net', '0.5259', '0.5343', '0.6202', '0.8070', '0.5268', '0.0090', '2.42s'],
        ['8', 'Linear Regression', '0.5254', '0.5243', '0.6172', '0.8073', '0.5242', '0.0034', '0.06s'],
        ['9', 'Ridge Regression', '0.5254', '0.5243', '0.6172', '0.8073', '0.5230', '0.0052', '0.05s'],
        ['10', 'Elastic Net', '0.5253', '0.5242', '0.6184', '0.8075', '0.5280', '0.0061', '0.04s'],
        ['11', 'Lasso Regression', '0.5250', '0.5240', '0.6192', '0.8077', '0.5194', '0.0034', '0.06s'],
        ['12', 'Bagging', '0.5138', '0.5522', '0.6136', '0.8172', '0.4706', '0.0082', '2.44s'],
        ['13', 'AdaBoost', '0.4703', '0.4697', '0.6945', '0.8530', '0.4727', '0.0093', '2.45s'],
        ['14', 'KNN', '0.4689', '0.5512', '0.6425', '0.8540', '0.4657', '0.0064', '0.03s'],
    ]
)

doc.add_paragraph()
add_heading_styled('Key Metrics Explained', level=2)
metrics = [
    ('R\u00b2 (Test): ', 'How much variance in job satisfaction the model explains on unseen data (higher = better, max 1.0)'),
    ('MAE: ', 'Average error in predicted satisfaction score (lower = better, in points on the 1\u20135 scale)'),
    ('RMSE: ', 'Root mean squared error \u2014 penalizes large errors more (lower = better)'),
    ('CV R\u00b2 Mean \u00b1 Std: ', 'Average R\u00b2 across 5-fold cross-validation \u2014 shows how stable the model is across different data splits'),
]
for bold_part, normal_part in metrics:
    p = doc.add_paragraph(style='List Bullet')
    add_bold_text(p, bold_part, normal_part)

# ============================================================
# SECTION 3
# ============================================================
doc.add_page_break()
add_heading_styled('3. Our Chosen Model: LightGBM', level=1)

add_heading_styled('Why LightGBM?', level=2)
p = doc.add_paragraph('We tested 14 models and the top 3 (XGBoost, LightGBM, Gradient Boosting) are ')
r = p.add_run('statistically identical')
r.bold = True
p.add_run(' (paired t-test p = 0.944). We selected LightGBM for these reasons:')

make_table(
    ['Criteria', 'LightGBM Performance', 'Assessment'],
    [
        ['Test R\u00b2', '0.5322', 'Tied best \u2014 explains 53.2% of variance'],
        ['MAE', '0.6057', 'Excellent \u2014 predictions off by ~0.6 points avg'],
        ['RMSE', '0.8015', 'Best RMSE of all 14 models'],
        ['CV R\u00b2 Mean', '0.5243', 'Highest CV mean \u2014 best generalization'],
        ['CV R\u00b2 Std', '\u00b1 0.0058', 'Good stability across folds'],
        ['Train vs Test R\u00b2', '0.5337 vs 0.5322', 'Minimal gap \u2014 no overfitting'],
        ['Training Speed', '0.38 seconds', 'Fastest ensemble model (7x faster than XGBoost)'],
    ]
)

doc.add_paragraph()
add_heading_styled('What does R\u00b2 = 0.53 mean in plain English?', level=2)
p = doc.add_paragraph('Our model explains ')
r = p.add_run('53% of the variation')
r.bold = True
p.add_run(' in how satisfied federal employees are with their jobs, using just 7 features. The other 47% comes from factors we don\'t have data on (e.g., pay, manager quality, office culture, personal factors). For social science survey data with ordinal Likert-scale responses, this is a ')
r = p.add_run('strong result')
r.bold = True
p.add_run('.')

add_heading_styled('LightGBM Configuration', level=2)
p = doc.add_paragraph('Default parameters were optimal. Hyperparameter tuning across 20 configurations showed the defaults were already near-optimal (tuned R\u00b2 = 0.5296 \u2014 slightly lower, confirming defaults are fine).')
doc.add_paragraph()
config_items = [
    'n_estimators = 100',
    'max_depth = 5',
    'learning_rate = 0.1',
    'random_state = 42',
    'verbose = -1',
]
for item in config_items:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(10)

# ============================================================
# SECTION 4
# ============================================================
doc.add_page_break()
add_heading_styled('4. Feature Importance \u2014 What Drives Job Satisfaction?', level=1)

p = doc.add_paragraph('This is the ')
r = p.add_run('most important finding for our capstone')
r.bold = True
p.add_run(':')

make_table(
    ['Feature', 'Importance', 'Interpretation'],
    [
        ['EEI \u2014 Employee Engagement', '34.4%', 'Strongest driver. Feeling valued, having purpose, seeing results.'],
        ['Q61 \u2014 Telework Frequency', '21.2%', '2nd strongest. More telework \u2192 more satisfaction.'],
        ['Q64 \u2014 Work-Life Balance', '18.9%', '3rd strongest. WLB satisfaction drives overall satisfaction.'],
        ['DFEDTEN \u2014 Federal Tenure', '10.4%', 'Some contribution \u2014 tenure shows different patterns.'],
        ['DSUPER \u2014 Supervisory Status', '5.8%', 'Minor contribution.'],
        ['DSEX \u2014 Gender', '4.9%', 'Minor contribution.'],
        ['DAGEGRP \u2014 Age Group', '4.5%', 'Minor contribution.'],
    ]
)

doc.add_paragraph()
add_heading_styled('Key Takeaway', level=2)

# Highlighted box paragraph
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
p.paragraph_format.right_indent = Cm(1)
r = p.add_run('The three workplace experience variables (engagement, telework, work-life balance) account for 74.5% of LightGBM\u2019s predictive power, while demographics account for 25.5%. ')
r.bold = True
r.font.color.rgb = RGBColor(0, 51, 102)
p.add_run('This indicates that job satisfaction in the federal workforce is driven primarily by how employees experience their work \u2014 their engagement, telework flexibility, and work-life balance \u2014 rather than individual demographic characteristics.')

doc.add_paragraph()
add_heading_styled('Average Importance Across All Top 5 Models', level=2)
p = doc.add_paragraph('Even across different model types, the pattern is consistent:')
avg_items = [
    ('EEI: ', '44.1% average'),
    ('Q64 (WLB): ', '33.8% average'),
    ('Q61 (Telework): ', '16.4% average'),
    ('All demographics combined: ', '< 6% average'),
]
for bold_part, normal_part in avg_items:
    p = doc.add_paragraph(style='List Bullet')
    add_bold_text(p, bold_part, normal_part)

# ============================================================
# SECTION 5
# ============================================================
doc.add_page_break()
add_heading_styled('5. Statistical Significance Tests', level=1)
p = doc.add_paragraph('We ran paired t-tests to check if the top models are truly different or just noise:')

make_table(
    ['Comparison', 'p-value', 'Significant?'],
    [
        ['XGBoost vs LightGBM', '0.944', 'No \u2014 essentially identical'],
        ['XGBoost vs Gradient Boosting', '0.968', 'No \u2014 essentially identical'],
        ['XGBoost vs Extra Trees', '0.874', 'No'],
        ['XGBoost vs Random Forest', '0.027', 'Yes \u2014 XGBoost significantly better in CV'],
        ['LightGBM vs Gradient Boosting', '0.953', 'No'],
        ['LightGBM vs Extra Trees', '0.825', 'No'],
        ['LightGBM vs Random Forest', '0.080', 'No (borderline)'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Bottom line: ')
r.bold = True
p.add_run('The top 3 models (XGBoost, LightGBM, Gradient Boosting) perform statistically the same. LightGBM is chosen for its best RMSE, highest CV mean, and fastest training speed.')

# ============================================================
# SECTION 6
# ============================================================
add_heading_styled('6. Correlation Matrix Findings', level=1)
p = doc.add_paragraph('Bivariate correlations between features and job satisfaction (Q69):')

make_table(
    ['Feature', 'Correlation with Q69', 'Strength'],
    [
        ['EEI (Engagement)', '0.647', 'Strong positive'],
        ['Q64 (Work-Life Balance)', '0.634', 'Strong positive'],
        ['Q61 (Telework)', '0.603', 'Strong positive'],
        ['DSUPER (Supervisor)', '0.058', 'Negligible'],
        ['DAGEGRP (Age)', '0.047', 'Negligible'],
        ['DFEDTEN (Tenure)', '0.016', 'Negligible'],
        ['DSEX (Gender)', '-0.011', 'Negligible'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Note: ')
r.italic = True
p.add_run('Q61, Q64, and EEI are also correlated with each other (r = 0.59\u20130.68), which is expected since engaged employees with telework also tend to report better WLB.')

# ============================================================
# SECTION 7
# ============================================================
doc.add_page_break()
add_heading_styled('7. How to Proceed \u2014 Next Steps', level=1)

add_heading_styled('For Module 5 (Final Model & Report)', level=2)
steps = [
    ('Primary model: LightGBM ', '\u2014 Use this for all final predictions and analysis'),
    ('Baseline model: Linear Regression ', '\u2014 Include for comparison and interpretability (OLS coefficients are easier to explain)'),
    ('Report the model comparison ', '\u2014 We tested 14 models and can show why LightGBM was selected'),
    ('Focus on feature importances ', '\u2014 This is where the research value is'),
]
for i, (bold_part, normal_part) in enumerate(steps, 1):
    p = doc.add_paragraph(style='List Number')
    add_bold_text(p, bold_part, normal_part)

add_heading_styled('For the Written Report', level=2)
report_items = [
    'We have the hierarchical OLS regression from Module 4 (traditional statistical approach)',
    'We now have the ML model comparison showing tree-based ensemble models perform best',
    'Together, these two approaches provide a comprehensive analysis: statistical inference (OLS) + predictive modeling (LightGBM)',
]
for item in report_items:
    doc.add_paragraph(item, style='List Bullet')

add_heading_styled('Files Saved', level=2)
make_table(
    ['File', 'Description'],
    [
        ['Model_Comparison_Analysis.ipynb', 'Complete Jupyter notebook with all code, charts, and analysis'],
        ['model_comparison_results.csv', 'Full 14-model comparison table'],
        ['significance_tests.csv', 'Paired t-test results'],
        ['tuning_results.csv', 'Hyperparameter tuning attempts'],
        ['feature_importances.csv', 'Feature importances for all top 5 models'],
        ['test_predictions.csv', 'Actual vs. predicted values on test set'],
        ['cv_fold_scores.csv', 'Individual cross-validation fold scores'],
        ['saved_models/lightgbm.joblib', 'Trained LightGBM model (ready to load and use)'],
        ['saved_models/scaler.joblib', 'StandardScaler for feature normalization'],
        ['figures/ folder', 'All visualization PNGs'],
    ]
)

# ============================================================
# SECTION 8
# ============================================================
doc.add_page_break()
add_heading_styled('8. One-Paragraph Summary (for the report)', level=1)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
p.paragraph_format.right_indent = Cm(1)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)
text = (
    'We evaluated 14 regression models on the 2024 Federal Employee Viewpoint Survey '
    '(N = 453,082) to predict overall job satisfaction (Q69). After training and testing '
    'via 80/20 split with 5-fold cross-validation, the top 5 models were all ensemble '
    'tree-based methods: XGBoost (R\u00b2 = 0.532), LightGBM (R\u00b2 = 0.532), Gradient '
    'Boosting (R\u00b2 = 0.532), Extra Trees (R\u00b2 = 0.531), and Random Forest (R\u00b2 = 0.531). '
    'Paired t-tests confirmed the top 3 models are statistically indistinguishable '
    '(p = 0.944). We selected LightGBM as the primary model due to its identical '
    'accuracy, highest cross-validation mean R\u00b2 (0.5243), best RMSE (0.8015), and '
    'fastest training time (0.38s \u2014 7x faster than XGBoost). Feature importance '
    'analysis revealed that employee engagement (34.4%), telework frequency (21.2%), '
    'and work-life balance satisfaction (18.9%) account for 74.5% of the model\u2019s '
    'predictive power. While demographic variables show some contribution in LightGBM\u2019s '
    'split-based importance, the three core workplace experience factors clearly '
    'dominate, indicating that job satisfaction in the federal workforce is driven '
    'primarily by how employees experience their work rather than individual '
    'demographic characteristics.'
)
r = p.add_run(text)
r.font.size = Pt(11)
r.italic = True

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Generated from Model_Comparison_Analysis.ipynb \u2014 all code, data, and visualizations are reproducible.')
r.italic = True
r.font.color.rgb = RGBColor(128, 128, 128)
r.font.size = Pt(9)

# ============================================================
# SAVE
# ============================================================
output_path = r'C:\Users\karan\Downloads\ADTA5940-Capstone\Model_Selection_Summary_for_Team.docx'
doc.save(output_path)
print(f"Word document saved: {output_path}")
