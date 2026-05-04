"""
Convert 14 Mendeley parenthetical citations to narrative format
and fix the missing em-dash in paragraph 23.
"""
import shutil
import base64
import json
import re
from docx import Document
from lxml import etree

NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

# --- Define which citations to convert ---
# Key: (paragraph_index, parenthetical_text) → narrative_text
CONVERSIONS = {
    (35, "(Greenhaus & Allen, 2011)"): "Greenhaus and Allen (2011)",
    (36, "(Shockley & Allen, 2007)"): "Shockley and Allen (2007)",
    (37, "(Fernandez et al., 2021)"): "Fernandez et al. (2021)",
    (38, "(Vaziri et al., 2020)"): "Vaziri et al. (2020)",
    (42, "(Harter et al., 2002)"): "Harter et al. (2002)",
    (43, "(Lavigna, 2013)"): "Lavigna (2013)",
    (43, "(Saks, 2006)"): "Saks (2006)",
    (44, "(Caillier, 2012)"): "Caillier (2012)",
    (44, "(Caillier, 2016)"): "Caillier (2016)",
    (47, "(Morganson et al., 2010)"): "Morganson et al. (2010)",
    (47, "(Golden & Fromen, 2011)"): "Golden and Fromen (2011)",
    (48, "(Lott & Chung, 2016)"): "Lott and Chung (2016)",
    (50, "(Bloom, 2024)"): "Bloom\u2019s (2024)",
    (51, "(Kaduk et al., 2019)"): "Kaduk et al. (2019)",
}

INPUT_FILE = "Module3_Scholarly_Review_Parekh_Lee_BACKUP.docx"
OUTPUT_FILE = "Module3_Scholarly_Review_Parekh_Lee_FIXED.docx"
BACKUP_FILE = "Module3_Scholarly_Review_Parekh_Lee_BACKUP.docx"


def get_sdt_display_text(sdt):
    """Get the visible text from an SDT element."""
    content = sdt.find('.//w:sdtContent', NS)
    if content is None:
        return ""
    parts = []
    for t in content.findall('.//w:t', NS):
        parts.append(t.text or "")
    return "".join(parts)


def update_mendeley_tag(sdt, new_display_text):
    """Update the Mendeley JSON in the tag to reflect manual override."""
    tag_el = sdt.find('.//w:tag', NS)
    if tag_el is None:
        return False

    tag_val = tag_el.get(f'{W}val', '')
    if not tag_val.startswith('MENDELEY_CITATION_v3_'):
        return False

    # Decode the base64 JSON
    b64_part = tag_val[len('MENDELEY_CITATION_v3_'):]
    try:
        data = json.loads(base64.b64decode(b64_part))
    except Exception:
        return False

    # Update the JSON to mark as manually overridden
    data['isEdited'] = True
    data['manualOverride']['isManuallyOverridden'] = True
    data['manualOverride']['manualOverrideText'] = new_display_text

    # Re-encode
    new_b64 = base64.b64encode(json.dumps(data).encode('utf-8')).decode('utf-8')
    new_tag = f'MENDELEY_CITATION_v3_{new_b64}'
    tag_el.set(f'{W}val', new_tag)
    return True


def update_sdt_display(sdt, new_text):
    """Replace the visible text inside the SDT content."""
    content = sdt.find('.//w:sdtContent', NS)
    if content is None:
        return False

    # Find the first run with text
    runs = content.findall('.//w:r', NS)
    if not runs:
        return False

    # Set the first run's text to the new narrative text
    first_t = runs[0].find('.//w:t', NS)
    if first_t is None:
        return False
    first_t.text = new_text
    # Preserve spacing
    first_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    # Remove text from any subsequent runs (shouldn't be any, but just in case)
    for r in runs[1:]:
        for t in r.findall('.//w:t', NS):
            t.text = ""

    return True


def fix_emdash(doc):
    """Fix 'at homereduced' → 'at home—reduced' in paragraph 23."""
    para = doc.paragraphs[23]
    for run in para.runs:
        if run.text and 'homereduced' in run.text:
            run.text = run.text.replace('homereduced', 'home\u2014reduced')
            print("  Fixed em-dash: 'homereduced' → 'home—reduced'")
            return True

    # If not found in a single run, check across consecutive runs
    runs = para.runs
    for i in range(len(runs) - 1):
        if runs[i].text and runs[i].text.rstrip().endswith('home'):
            next_text = runs[i + 1].text if runs[i + 1].text else ""
            if next_text.lstrip().startswith('reduced'):
                # Check if there's already an em-dash
                combined = runs[i].text + next_text
                if 'home' in combined and 'reduced' in combined and '\u2014' not in combined:
                    runs[i].text = runs[i].text.rstrip()
                    if runs[i].text.endswith('home'):
                        runs[i].text = runs[i].text + '\u2014'
                        runs[i + 1].text = runs[i + 1].text.lstrip()
                        if runs[i + 1].text.startswith('reduced'):
                            print("  Fixed em-dash across runs")
                            return True
    
    # broader search
    full = para.text
    if 'home' in full and 'reduced' in full and '\u2014' not in full:
        print(f"  WARNING: em-dash issue detected but couldn't fix automatically in P23")
        print(f"  Paragraph text around issue: ...{full[full.find('home')-5:full.find('reduced')+15]}...")
        return False
    elif '\u2014' in full or '—' in full:
        print("  Em-dash already present in P23")
        return True
    else:
        print("  Could not locate em-dash issue in P23")
        return False


def main():
    # Backup already exists from prior run
    print(f"Reading from {INPUT_FILE}")

    doc = Document(INPUT_FILE)
    changed = 0

    for para_idx, para in enumerate(doc.paragraphs):
        el = para._element
        sdts = el.findall('.//w:sdt', NS)

        for sdt in sdts:
            display = get_sdt_display_text(sdt)

            key = (para_idx, display)
            if key in CONVERSIONS:
                new_text = CONVERSIONS[key]
                print(f"P{para_idx}: '{display}' → '{new_text}'")

                # Update the Mendeley JSON metadata
                if update_mendeley_tag(sdt, new_text):
                    print("  ✓ Mendeley tag updated")
                else:
                    print("  ✗ Failed to update Mendeley tag")

                # Update the visible display text
                if update_sdt_display(sdt, new_text):
                    print("  ✓ Display text updated")
                    changed += 1
                else:
                    print("  ✗ Failed to update display text")

    print(f"\nConverted {changed} of {len(CONVERSIONS)} citations")

    # Fix em-dash
    print("\nFixing em-dash in P23...")
    fix_emdash(doc)

    # Save
    doc.save(OUTPUT_FILE)
    print(f"\nSaved to {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
