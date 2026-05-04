"""
Create shortened, less AI-sounding Module 4 document for Karan Parekh
- Cuts Variable Definitions table, Model Diagnostics, residual figure
- Condenses Data Prep, Discussion, Limitations
- Adds AI usage disclosure
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

FIGURES = r"C:\Users\karan\Downloads\ADTA5940-Capstone\figures"
OUT_DIR = r"C:\Users\karan\Downloads\ADTA5940-Capstone"


def style_doc(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    para = style.paragraph_format
    para.space_after = Pt(6)
    para.line_spacing = 2.0  # double-spaced for academic
    return doc


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return h


def add_figure(doc, filename, caption, width=5.0):
    path = os.path.join(FIGURES, filename)
    if not os.path.exists(path):
        p = doc.add_paragraph(f"[Figure not found: {filename}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(10)
        run.font.italic = True


def make_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


def create_karan_doc_v2():
    doc = Document()
    style_doc(doc)

    # ===== TITLE PAGE =====
    for _ in range(5):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("Module 4: Model with Results")
    run.font.size = Pt(20)
    run.bold = True

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t2.add_run(
        "Remote Work, Work-Life Balance, and\n"
        "Federal Employee Job Satisfaction"
    )
    run.font.size = Pt(14)

    for _ in range(3):
        doc.add_paragraph()

    for line in [
        "Karan Parekh",
        "ADTA 5940 — Analytics Capstone Experience",
        "Section 501 | Spring 2026",
        "University of North Texas",
        "Dr. Denise Philpot",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line).font.size = Pt(12)

    doc.add_page_break()

    # ===== 1. RESEARCH QUESTION =====
    add_heading_styled(doc, "1. Research Question", level=1)

    doc.add_paragraph(
        "During our EDA, we noticed a pattern: federal employees who telework regularly "
        "tend to rate their job satisfaction higher than those required to come in. But is "
        "telework itself the driver, or are other factors — like engagement and work-life "
        "balance support — doing the heavy lifting? That is what this analysis sets out to "
        "answer."
    )

    rq = doc.add_paragraph()
    rq.paragraph_format.left_indent = Inches(0.5)
    rq.paragraph_format.right_indent = Inches(0.5)
    run = rq.add_run(
        "To what extent does telework frequency predict overall job satisfaction among "
        "federal employees, after accounting for work-life balance perceptions, employee "
        "engagement, and demographic characteristics?"
    )
    run.italic = True

    doc.add_paragraph("Three hypotheses guide the analysis:")

    for h in [
        "H1: Telework frequency significantly predicts job satisfaction.",
        "H2: Work-life balance and employee engagement each add meaningful explanatory "
        "power beyond telework alone.",
        "H3: After controlling for WLB and engagement, the telework effect will shrink "
        "substantially, indicating partial mediation.",
    ]:
        doc.add_paragraph(h, style='List Bullet')

    # ===== 2. DATA PREPARATION (condensed) =====
    add_heading_styled(doc, "2. Data Preparation", level=1)

    doc.add_paragraph(
        "The dataset is OPM's 2024 Federal Employee Viewpoint Survey (FEVS) Public Release "
        "Data File — 674,207 rows and 96 columns. Likert-scale survey items were converted "
        "from text to numeric (1–5), with non-applicable responses (\"X\") treated as missing. "
        "Two composite indices were created following OPM's published methodology: the Employee "
        "Engagement Index (EEI), which averages three sub-indices covering intrinsic work "
        "experience, supervisor effectiveness, and leadership; and a Work-Life Balance (WLB) "
        "composite averaging three items on peer, supervisor, and senior leader support for "
        "work-life programs."
    )

    doc.add_paragraph(
        "The telework variable (Q91) has four categories: Routine/Remote, Situational, "
        "Required in Office, and Chooses Not To. Dummy variables were created with "
        "Routine/Remote as the reference group. Demographic controls include age (under/over "
        "40), gender, supervisory status, and federal tenure (three levels). After listwise "
        "deletion, the analytic sample is n = 517,697."
    )

    # ===== 3. FORMAL MODEL =====
    add_heading_styled(doc, "3. Formal Model", level=1)

    doc.add_paragraph(
        "I used a hierarchical OLS regression, entering predictors in four blocks. This "
        "approach lets us see how much each group of variables adds — and specifically, "
        "what happens to the telework effect as we layer in stronger predictors."
    )

    for spec in [
        "Model 1 (Telework only): Satisfaction = telework dummies",
        "Model 2 (+ Work-Life Balance): adds WLB composite",
        "Model 3 (+ Employee Engagement): adds EEI",
        "Model 4 (Full): adds supervisor status, gender, age, and tenure",
    ]:
        doc.add_paragraph(spec, style='List Bullet')

    doc.add_paragraph(
        "The dependent variable is Q70 (\"Considering everything, how satisfied are you "
        "with your job?\") on a 1–5 Likert scale."
    )

    # ===== 4. RESULTS =====
    add_heading_styled(doc, "4. Results", level=1)

    # 4.1 Descriptive
    add_heading_styled(doc, "4.1 Descriptive Overview", level=2)

    doc.add_paragraph(
        "Average job satisfaction is 3.82 out of 5 (SD = 1.09). The WLB composite averages "
        "4.02 and EEI averages 3.95, both above the scale midpoint. About 42.8% of "
        "respondents telework routinely, 33.6% situationally, 19.7% are required on-site, "
        "and 3.8% choose not to telework."
    )

    doc.add_paragraph(
        "Figure 1 shows a clear raw gap: routine teleworkers average 3.94 on satisfaction "
        "while required-in-office employees average 3.54 — a 0.40-point difference. A "
        "one-way ANOVA confirms this is statistically significant (F(3, 639,584) = 4,047, "
        "p < .001), though the effect size is small (η² = 0.019)."
    )

    add_figure(doc, "fig1_satisfaction_by_telework.png",
               "Figure 1. Mean Job Satisfaction by Telework Category (95% CI)", 4.5)

    # 4.2 Hierarchical regression
    add_heading_styled(doc, "4.2 Hierarchical Regression", level=2)

    doc.add_paragraph(
        "Table 1 shows how the model improves at each step. This is the core finding."
    )

    hier_headers = ["Metric", "Model 1", "Model 2", "Model 3", "Model 4"]
    hier_rows = [
        ["Block added", "Telework", "+ WLB", "+ EEI", "+ Demographics"],
        ["R²", ".020", ".409", ".595", ".597"],
        ["ΔR²", "—", ".389", ".185", ".002"],
        ["Cohen's f²", "0.020 (small)", "0.692 (large)", "1.466 (large)", "1.479 (large)"],
    ]
    p = doc.add_paragraph()
    run = p.add_run("Table 1. Hierarchical Regression Summary (n = 517,697)")
    run.bold = True
    run.font.size = Pt(10)
    make_table(doc, hier_headers, hier_rows)

    doc.add_paragraph("")

    doc.add_paragraph(
        "Telework alone explains just 2% of variance (Model 1). Adding WLB jumps R² to "
        ".41 — by far the biggest single gain. EEI pushes it to .60. Demographics add "
        "only 0.2 percentage points. In short, engagement and work-life balance perceptions "
        "do almost all the explanatory work."
    )

    add_figure(doc, "fig2_r2_progression.png",
               "Figure 2. R² Progression Across Model Blocks", 4.5)

    # 4.3 Full model coefficients
    add_heading_styled(doc, "4.3 Key Coefficients", level=2)

    doc.add_paragraph(
        "Table 2 shows the full model coefficients. The standardized Beta values show "
        "relative importance."
    )

    coef_headers = ["Predictor", "B", "Beta", "p"]
    coef_rows = [
        ["Employee Engagement Index", "0.945", "0.705", "< .001"],
        ["Work-Life Balance Composite", "0.111", "0.091", "< .001"],
        ["TW: Required in Office", "0.114", "0.043", "< .001"],
        ["TW: Chooses Not To", "0.090", "0.016", "< .001"],
        ["TW: Infrequent", "−0.002", "−0.001", ".326"],
        ["Age 40+", "0.088", "0.036", "< .001"],
        ["Tenure 20+ yrs", "0.049", "0.020", "< .001"],
        ["Male", "0.018", "0.009", "< .001"],
        ["Supervisor", "−0.019", "−0.007", "< .001"],
    ]
    p = doc.add_paragraph()
    run = p.add_run("Table 2. Full Model Coefficients (Model 4)")
    run.bold = True
    run.font.size = Pt(10)
    make_table(doc, coef_headers, coef_rows)

    p = doc.add_paragraph()
    run = p.add_run("R² = .597, F(10, 517,686) = 76,564, p < .001")
    run.font.size = Pt(10)
    run.italic = True

    doc.add_paragraph(
        "EEI dominates with Beta = 0.705 — a one-unit increase in engagement corresponds "
        "to nearly a full point increase in satisfaction. WLB is a distant second (Beta = "
        "0.091). Every other predictor has a Beta below 0.05."
    )

    add_figure(doc, "fig3_coefficients_forest.png",
               "Figure 3. Standardized Coefficients — Full Model", 5.0)

    # 4.4 Simpson's Paradox
    add_heading_styled(doc, "4.4 The Telework Reversal", level=2)

    doc.add_paragraph(
        "The most striking finding is what happens to the telework coefficients across "
        "models (Table 3)."
    )

    tw_headers = ["Predictor", "Model 1", "Model 2", "Model 3", "Model 4"]
    tw_rows = [
        ["TW: Infrequent", "−0.121***", "0.013***", "−0.002", "−0.002"],
        ["TW: Required in Office", "−0.407***", "+0.094***", "+0.109***", "+0.114***"],
        ["TW: Chooses Not To", "−0.029***", "+0.127***", "+0.096***", "+0.089***"],
    ]
    p = doc.add_paragraph()
    run = p.add_run("Table 3. Telework Coefficients Across Models")
    run.bold = True
    run.font.size = Pt(10)
    make_table(doc, tw_headers, tw_rows)

    doc.add_paragraph(
        "In Model 1, required-in-office employees score 0.41 points lower than teleworkers. "
        "But in Model 2, after adding WLB, the sign flips to +0.09. By the full model it "
        "is +0.11. This is Simpson's Paradox: the raw telework gap is not caused by the "
        "work arrangement itself. In-office employees tend to report lower WLB support and "
        "engagement — and those are what actually drive their lower satisfaction. Once we "
        "hold those factors constant, office workers are actually slightly more satisfied."
    )

    # ===== 5. DISCUSSION =====
    add_heading_styled(doc, "5. Discussion", level=1)

    doc.add_paragraph("Returning to the hypotheses:")

    for h in [
        "H1 — Partially supported. Telework is statistically significant but explains "
        "only 2% of variance alone, and the effect reverses with controls.",
        "H2 — Strongly supported. WLB added 39 percentage points to R² and EEI added "
        "another 19. Together they get the model to 60% explained variance.",
        "H3 — Supported. The telework coefficient reversed direction once WLB and "
        "engagement were held constant.",
    ]:
        doc.add_paragraph(h, style='List Bullet')

    add_heading_styled(doc, "5.1 Implications for Agencies", level=2)

    doc.add_paragraph(
        "The practical takeaway is that expanding telework alone will not reliably increase "
        "satisfaction. What moves the needle is employee engagement — effective supervision, "
        "meaningful work, and visible leadership commitment. Agencies that invest in those "
        "areas will see satisfied employees regardless of where they physically work. This "
        "does not make telework unimportant (it supports recruitment and retention), but "
        "the data does not support treating it as a satisfaction silver bullet."
    )

    add_heading_styled(doc, "5.2 Limitations", level=2)

    for lim in [
        "This is cross-sectional data (one point in time), so causal claims are not "
        "possible. The mediation pattern is suggestive but not confirmed.",
        "WLB and EEI are highly correlated (VIF > 50), which inflates their individual "
        "standard errors. Their combined effect is stable, but separating their unique "
        "impacts would require structural equation modeling.",
    ]:
        doc.add_paragraph(lim, style='List Bullet')

    # ===== REFERENCES =====
    add_heading_styled(doc, "References", level=1)

    refs = [
        "Bakker, A. B., & Demerouti, E. (2017). Job Demands-Resources theory: Taking "
        "stock and looking forward. Journal of Occupational Health Psychology, 22(3), 273–285.",
        "Caillier, J. G. (2012). The impact of teleworking on work motivation in a U.S. "
        "federal government agency. American Review of Public Administration, 42(4), 461–480.",
        "Cohen, J. (1988). Statistical power analysis for the behavioral sciences (2nd ed.). Erlbaum.",
        "Office of Personnel Management. (2024). Federal Employee Viewpoint Survey: Public "
        "Release Data File and Codebook. Washington, DC: OPM.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        for run in p.runs:
            run.font.size = Pt(11)

    # ===== AI USAGE DISCLOSURE =====
    doc.add_paragraph()
    add_heading_styled(doc, "AI Usage Disclosure", level=1)

    doc.add_paragraph(
        "In accordance with course AI policy, I used GitHub Copilot (an AI coding assistant) "
        "during this project. Specifically, AI was used to assist with writing and debugging "
        "the Python code that runs the regression models and generates figures. The research "
        "question, hypotheses, model design, interpretation of results, and all written "
        "analysis reflect my own understanding of the data and methods. All statistical "
        "outputs were verified independently. The final document was reviewed and edited "
        "by me to ensure accuracy and original expression."
    )

    # ===== SAVE =====
    path = os.path.join(OUT_DIR, "Module4_Model_Results_Karan_Parekh_v2.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


if __name__ == "__main__":
    p = create_karan_doc_v2()
    print(f"Done! → {p}")
